# coding: utf-8

from dataclasses import dataclass
from datetime import datetime, timezone
import importlib.util
from pathlib import Path
import queue
import shutil
import struct
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import traceback
import zipfile
from xml.sax.saxutils import escape

import numpy as np

from qar_common import numeric_array, numeric_columns, read_qar_csv

ROOT_DIR = Path(__file__).resolve().parent
PICS_DIR = ROOT_DIR / "pics"
REPORT_BASENAME = "QAR分析"
AIRCRAFT_ORDER = ("CFM", "LEAP", "PW")
AIRCRAFT_DISPLAY = {
    "CFM": "CFM",
    "LEAP": "LEAP",
    "PW": "PW",
}


@dataclass
class AnalysisTask:
    person: str
    aircraft: str
    paths: list
    output_path: Path


AIRCRAFT_DIRS = {
    "cfm": "CFM",
    "leap": "LEAP",
    "pw": "PW",
}

MODULE_PATHS = {
    "CFM": ROOT_DIR / "test_cfm" / "pic_code.py",
    "LEAP": ROOT_DIR / "test_leap" / "pic_5code.py",
    "PW": ROOT_DIR / "test_pw" / "pic_code.py",
}


def classify_csv(path):
    name = path.name.lower()
    has_qlzlly = "qlzlly" in name
    has_qlz = "qlz" in name
    has_lly = "lly" in name

    if has_qlzlly:
        return "LEAP"
    if has_qlz:
        return "CFM"
    if has_lly:
        return "PW"
    return None


def unique_output_path(person, aircraft):
    PICS_DIR.mkdir(parents=True, exist_ok=True)
    base = PICS_DIR / f"{person}_{aircraft}.png"
    if not base.exists():
        return base

    index = 1
    while True:
        candidate = PICS_DIR / f"{person}_{aircraft}_{index}.png"
        if not candidate.exists():
            return candidate
        index += 1


def unique_report_path():
    return ROOT_DIR / f"{REPORT_BASENAME}.docx"


def scan_data_root(data_root):
    tasks = []
    skipped = []
    errors = []

    for person_dir in sorted((p for p in data_root.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        aircraft_dirs = {}
        for child in person_dir.iterdir():
            if child.is_dir() and child.name.lower() in AIRCRAFT_DIRS:
                aircraft_dirs[child.name.lower()] = child

        for dir_key in ("cfm", "leap", "pw"):
            aircraft = AIRCRAFT_DIRS[dir_key]
            aircraft_dir = aircraft_dirs.get(dir_key)
            if aircraft_dir is None:
                skipped.append((person_dir.name, aircraft, "缺少机型文件夹"))
                continue

            csv_paths = sorted(aircraft_dir.glob("*.csv"), key=lambda p: p.name.lower())
            if not csv_paths:
                skipped.append((person_dir.name, aircraft, "没有 CSV 文件"))
                continue

            wrong_paths = [path for path in csv_paths if classify_csv(path) != aircraft]
            if wrong_paths:
                bad_names = "、".join(path.name for path in wrong_paths[:5])
                if len(wrong_paths) > 5:
                    bad_names += f" 等 {len(wrong_paths)} 个文件"
                errors.append((person_dir.name, aircraft, f"{person_dir.name} 的 {aircraft} 数据错误：{bad_names}"))
                continue

            tasks.append(
                AnalysisTask(
                    person=person_dir.name,
                    aircraft=aircraft,
                    paths=csv_paths,
                    output_path=unique_output_path(person_dir.name, aircraft),
                )
            )

    return tasks, skipped, errors


def grouped_tasks(tasks):
    grouped = {}
    for task in tasks:
        grouped.setdefault(task.person, {})[task.aircraft] = task
    return dict(sorted(grouped.items(), key=lambda item: item[0].lower()))


def load_analysis_modules():
    modules = {}
    for aircraft, module_path in MODULE_PATHS.items():
        spec = importlib.util.spec_from_file_location(f"qar_{aircraft.lower()}_analysis", module_path)
        if spec is None or spec.loader is None:
            raise RuntimeError(f"无法加载 {aircraft} 分析模块：{module_path}")
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        modules[aircraft] = module
    return modules


def _collect_cfm_takeoff_points(path):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    dy = "N1 Actual Eng 1"
    height = "RADIO HEIGHT (R/A1) SYS. 1_1"
    pitch_cols = [
        "Capt Pitch Command Positi_1",
        "Capt Pitch Command Positi_2",
        "Capt Pitch Command Positi_3",
        "Capt Pitch Command Positi_4",
        "Capt Pitch Command Positi_5",
        "Capt Pitch Command Positi_6",
        "Capt Pitch Command Positi_7",
        "Capt Pitch Command Positi_8",
    ]
    roll_cols = [
        "Capt Roll Command Positio_1",
        "Capt Roll Command Positio_2",
        "Capt Roll Command Positio_3",
        "Capt Roll Command Positio_4",
        "Capt Roll Command Positio_5",
        "Capt Roll Command Positio_6",
        "Capt Roll Command Positio_7",
        "Capt Roll Command Positio_8",
    ]
    required = [dy, height] + pitch_cols + roll_cols
    if any(column not in data.columns for column in required):
        return np.array([]), np.array([])

    heights = numeric_columns(data, [height]).iloc[:, 0]
    n1 = numeric_columns(data, [dy]).iloc[:, 0]
    selected = data.loc[(heights < 50) & (n1 >= 50)]
    if selected.empty:
        return np.array([]), np.array([])

    selected_indexes = list(selected.index)
    cut_count = len(selected_indexes)
    for index in range(1, len(selected_indexes)):
        if selected_indexes[index] - selected_indexes[index - 1] > 2:
            cut_count = max(index - 1, 0)
            break
    selected = selected.loc[selected_indexes[:cut_count]]
    if selected.empty:
        return np.array([]), np.array([])

    pitch = numeric_columns(selected, pitch_cols)
    selected = selected.loc[(pitch >= 0).sum(axis=1) == 0]
    if selected.empty:
        return np.array([]), np.array([])

    x_values = numeric_array(selected, roll_cols).reshape(-1)
    y_values = -numeric_array(selected, pitch_cols).reshape(-1)
    return x_values[np.isfinite(x_values)], y_values[np.isfinite(y_values)]


def _collect_pw_takeoff_points(path, leap=False):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    pitch_cols = [
        "PITCH CAPT CMD POSITION_58",
        "PITCH CAPT CMD POSITION_186",
        "PITCH CAPT CMD POSITION_314",
        "PITCH CAPT CMD POSITION_442",
        "PITCH CAPT CMD POSITION_570",
        "PITCH CAPT CMD POSITION_698",
        "PITCH CAPT CMD POSITION_826",
        "PITCH CAPT CMD POSITION_954",
    ]
    roll_cols = [
        "ROLL CAPT CMD POSITION_46",
        "ROLL CAPT CMD POSITION_174",
        "ROLL CAPT CMD POSITION_302",
        "ROLL CAPT CMD POSITION_430",
        "ROLL CAPT CMD POSITION_558",
        "ROLL CAPT CMD POSITION_686",
        "ROLL CAPT CMD POSITION_814",
        "ROLL CAPT CMD POSITION_942",
    ]
    n1_col = "N1 ACTUAL (LOW ROTOR SPEED) SYS. 1" if leap else "N1 TARGET SYS 1"
    height_col = "RADIO ALTITUDE SYS. 1_5" if leap else "RADIO HEIGHT (R/A1) SYS. 1_5"
    required = [n1_col, height_col] + pitch_cols + roll_cols
    if any(column not in data.columns for column in required):
        return np.array([]), np.array([])

    n1 = numeric_columns(data, [n1_col]).iloc[:, 0]
    selected = data.loc[n1 >= 50]
    if selected.empty:
        return np.array([]), np.array([])

    pitch = numeric_columns(selected, pitch_cols)
    selected = selected.loc[(pitch < 0).sum(axis=1) > 0]
    if selected.empty:
        return np.array([]), np.array([])

    heights = numeric_columns(selected, [height_col]).iloc[:, 0]
    indexes = []
    for row_index, value in heights.items():
        indexes.append(row_index)
        if value > 50:
            break
    selected = data.loc[indexes]
    if selected.empty:
        return np.array([]), np.array([])

    x_values = numeric_array(selected, roll_cols).reshape(-1)
    y_values = -numeric_array(selected, pitch_cols).reshape(-1)
    if leap:
        x_values = -x_values
    return x_values[np.isfinite(x_values)], y_values[np.isfinite(y_values)]


def collect_takeoff_points(task):
    collectors = {
        "CFM": lambda path: _collect_cfm_takeoff_points(path),
        "LEAP": lambda path: _collect_pw_takeoff_points(path, leap=True),
        "PW": lambda path: _collect_pw_takeoff_points(path, leap=False),
    }
    x_parts = []
    y_parts = []
    for path in task.paths:
        try:
            x_values, y_values = collectors[task.aircraft](path)
        except Exception:
            continue
        if len(x_values) and len(y_values):
            x_parts.append(x_values)
            y_parts.append(y_values)
    if not x_parts or not y_parts:
        return np.array([]), np.array([])
    return np.concatenate(x_parts), np.concatenate(y_parts)


def finite_values(values):
    values = np.asarray(values, dtype=float).reshape(-1)
    return values[np.isfinite(values)]


def has_columns(dataframe, columns):
    return all(column in dataframe.columns for column in columns)


def range_indexes_by_height(dataframe, height_columns, lower, upper):
    max_height = numeric_columns(dataframe, height_columns).max(axis=1)
    indexes = []
    for index in reversed(max_height.index):
        value = max_height.loc[index]
        if lower < value < upper:
            indexes.append(index)
        elif value > upper:
            break
    return indexes


def control_metrics(x_values, y_values, wind_values=None):
    x_values = finite_values(x_values)
    y_values = finite_values(y_values)
    if not len(x_values) or not len(y_values):
        return None

    below_axis = y_values[y_values < 0]
    wind_values = finite_values(wind_values if wind_values is not None else [])
    return {
        "pull": max(float(np.nanmax(y_values)), 0.0),
        "steady": abs(float(np.nanmin(below_axis))) if len(below_axis) else 0.0,
        "lateral": float(np.nanmax(np.abs(x_values))),
        "wind": float(np.nanmax(wind_values)) if len(wind_values) else None,
    }


def merge_control_metrics(metric_parts, require_wind=False):
    metric_parts = [metrics for metrics in metric_parts if metrics]
    if not metric_parts:
        return None

    wind_values = [metrics["wind"] for metrics in metric_parts if metrics.get("wind") is not None]
    if require_wind and not wind_values:
        return None

    return {
        "pull": max(metrics["pull"] for metrics in metric_parts),
        "steady": max(metrics["steady"] for metrics in metric_parts),
        "lateral": max(metrics["lateral"] for metrics in metric_parts),
        "wind": max(wind_values) if wind_values else None,
    }


def format_number(value):
    value = float(value)
    if value >= 0:
        return str(int(np.floor(value + 0.5)))
    return str(int(np.ceil(value - 0.5)))


def format_value_range(values, unit=""):
    values = finite_values(values)
    if not len(values):
        return None

    low = float(np.nanmin(values))
    high = float(np.nanmax(values))
    if abs(high - low) < 0.05:
        return f"{format_number(low)}{unit}"
    return f"{format_number(low)}~{format_number(high)}{unit}"


ALERT_THRESHOLD = 10


def analysis_part(text, alert=False):
    return {"text": text, "alert": bool(alert)}


def metric_analysis_part(prefix, value, suffix):
    displayed_value = format_number(value)
    return analysis_part(f"{prefix}{displayed_value}{suffix}", int(displayed_value) > ALERT_THRESHOLD)


def strip_trailing_period(parts):
    copied = [dict(part) for part in parts if part.get("text")]
    if copied and copied[-1]["text"].endswith("。"):
        copied[-1]["text"] = copied[-1]["text"][:-1]
    return [part for part in copied if part["text"]]


def join_analysis_parts(parts_iter):
    joined = []
    has_content = False
    for parts in parts_iter:
        cleaned = strip_trailing_period(parts)
        if not cleaned:
            continue
        if has_content:
            joined.append(analysis_part("；"))
        joined.extend(cleaned)
        has_content = True
    if not has_content:
        return [analysis_part("该项数据无法自动计算。")]
    joined.append(analysis_part("。"))
    return joined


def takeoff_analysis_parts(task):
    x_values, y_values = collect_takeoff_points(task)
    metrics = control_metrics(x_values, y_values)
    if metrics is None:
        return [analysis_part(f"{task.aircraft}机型起飞50ft以下杆量无法自动计算。")]

    if metrics["steady"] <= 0:
        status = "起飞杆量正常"
        steady_parts = []
    elif metrics["steady"] <= 5:
        status = "起飞有少量稳杆"
        steady_parts = [
            analysis_part("，"),
            metric_analysis_part("稳杆量最大约", metrics["steady"], "个单位"),
        ]
    else:
        status = "起飞有稳杆"
        steady_parts = [
            analysis_part("，"),
            metric_analysis_part("稳杆量最大约", metrics["steady"], "个单位"),
        ]

    return [
        analysis_part(f"{task.aircraft}机型{status}，"),
        metric_analysis_part("带杆量最大约", metrics["pull"], "个单位"),
        *steady_parts,
        analysis_part("，"),
        metric_analysis_part("横侧杆量", metrics["lateral"], "个单位"),
        analysis_part("。"),
    ]


def takeoff_analysis_text(task):
    return "".join(part["text"] for part in takeoff_analysis_parts(task))


def control_metrics_alert(metrics, include_wind=False):
    if not metrics:
        return False
    if any(metrics[key] > ALERT_THRESHOLD for key in ("pull", "steady", "lateral")):
        return True
    return bool(include_wind and metrics.get("wind") is not None and metrics["wind"] > ALERT_THRESHOLD)


def takeoff_analysis_alert(task):
    x_values, y_values = collect_takeoff_points(task)
    return control_metrics_alert(control_metrics(x_values, y_values))


def _collect_cfm_liftoff_values(path):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    aircraft_state_cols = [
        "AIRCRAFT ON GROUND (CFM)_592",
        "AIRCRAFT ON GROUND (CFM)_599",
    ]
    gear_cols = [
        "LANDING GEAR NOSE_1",
        "LANDING GEAR NOSE_2",
        "LANDING GEAR NOSE_3",
        "LANDING GEAR NOSE_4",
    ]
    pitch_cols = [
        "Pitch attitude CA_1",
        "Pitch attitude CA_2",
        "Pitch attitude CA_3",
        "Pitch attitude CA_4",
        "Pitch attitude CA_5",
        "Pitch attitude CA_6",
        "Pitch attitude CA_7",
    ]
    height_cols = [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
        "RADIO HEIGHT (R/A1) SYS. 2_1",
        "RADIO HEIGHT (R/A1) SYS. 2_2",
        "RADIO HEIGHT (R/A1) SYS. 2_3",
        "RADIO HEIGHT (R/A1) SYS. 2_4",
    ]
    required = aircraft_state_cols + gear_cols + pitch_cols + height_cols
    if not has_columns(data, required):
        return np.array([])

    gear_air_counts = (data[gear_cols] == "AIR").sum(axis=1)
    gear_indexes = []
    for index, count in gear_air_counts.items():
        if count > 0:
            for offset in range(4):
                candidate = index - offset
                if candidate in data.index:
                    gear_indexes.append(candidate)

    max_height = numeric_columns(data, height_cols).max(axis=1)
    height_indexes = []
    for index, value in max_height.items():
        if value > 50:
            break
        height_indexes.append(index)

    candidate_indexes = sorted(set(gear_indexes).intersection(height_indexes))
    if not candidate_indexes:
        return np.array([])

    selected = data.loc[candidate_indexes]
    liftoff_mask = (selected[aircraft_state_cols[1]] == "NO").astype(int) - (
        selected[aircraft_state_cols[0]] == "YES"
    ).astype(int) > 0
    liftoff_indexes = list(liftoff_mask[liftoff_mask].index)
    if not liftoff_indexes:
        return np.array([])

    return finite_values(numeric_array(data.loc[[liftoff_indexes[0]]], pitch_cols))


def _collect_pw_liftoff_values(path, leap=False):
    if leap:
        return np.array([])

    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    gear_cols = [
        "NOSE LANDING GEAR COMPRESSED SYS. 1_49",
        "NOSE LANDING GEAR COMPRESSED SYS. 1_305",
        "NOSE LANDING GEAR COMPRESSED SYS. 1_561",
        "NOSE LANDING GEAR COMPRESSED SYS. 1_817",
    ]
    pitch_cols = [
        "PITCH ANGLE SYS. 1_43",
        "PITCH ANGLE SYS. 1_107",
        "PITCH ANGLE SYS. 1_171",
        "PITCH ANGLE SYS. 1_235",
    ]
    height_col = "RADIO HEIGHT (R/A1) SYS. 1_4"
    ground_cols = [
        "GROUND FLIGHT BOOLEAN BSOL_26",
        "GROUND FLIGHT BOOLEAN BSOL_282",
        "GROUND FLIGHT BOOLEAN BSOL_538",
        "GROUND FLIGHT BOOLEAN BSOL_794",
    ]
    required = gear_cols + pitch_cols + [height_col] + ground_cols
    if not has_columns(data, required):
        return np.array([])

    gear_mask = (data[gear_cols] == "NOFALT").sum(axis=1) > 0
    gear_positions = np.where(gear_mask)[0]
    if not len(gear_positions):
        return np.array([])
    start = max(int(gear_positions[0]) - 3, 0)

    height_values = numeric_columns(data, [height_col]).iloc[:, 0]
    height_positions = np.where(height_values > 50)[0]
    if not len(height_positions):
        return np.array([])
    end = int(height_positions[0])
    if end < start:
        return np.array([])

    selected = data.loc[start:end]
    liftoff_mask = (selected[ground_cols] == "NOACT").sum(axis=1) > 0
    liftoff_indexes = list(liftoff_mask[liftoff_mask].index)
    if not liftoff_indexes:
        return np.array([])

    return finite_values(numeric_array(data.loc[[liftoff_indexes[0]]], pitch_cols))


def collect_liftoff_values(task):
    if task.aircraft == "LEAP":
        return np.array([])

    collectors = {
        "CFM": _collect_cfm_liftoff_values,
        "PW": lambda path: _collect_pw_liftoff_values(path, leap=False),
    }
    values = []
    for path in task.paths:
        try:
            values_part = collectors[task.aircraft](path)
        except Exception:
            continue
        if len(values_part):
            values.append(values_part)
    if not values:
        return np.array([])
    return np.concatenate(values)


def liftoff_analysis_text(task):
    if task.aircraft == "LEAP":
        return "LEAP机型离地姿态数据图形有误。"

    values = collect_liftoff_values(task)
    value_range = format_value_range(values, "度")
    if value_range is None:
        return f"{task.aircraft}机型离地姿态数据无法自动计算。"
    return f"{task.aircraft}机型离地曲线正常，离地姿态范围{value_range}。"


def _collect_cfm_phase_metrics(path, lower, upper):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    height_cols = [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
        "RADIO HEIGHT (R/A1) SYS. 2_1",
        "RADIO HEIGHT (R/A1) SYS. 2_2",
        "RADIO HEIGHT (R/A1) SYS. 2_3",
        "RADIO HEIGHT (R/A1) SYS. 2_4",
    ]
    pitch_cols = [
        "Capt Pitch Command Positi_1",
        "Capt Pitch Command Positi_2",
        "Capt Pitch Command Positi_3",
        "Capt Pitch Command Positi_4",
        "Capt Pitch Command Positi_5",
        "Capt Pitch Command Positi_6",
        "Capt Pitch Command Positi_7",
        "Capt Pitch Command Positi_8",
    ]
    roll_cols = [
        "Capt Roll Command Positio_1",
        "Capt Roll Command Positio_2",
        "Capt Roll Command Positio_3",
        "Capt Roll Command Positio_4",
        "Capt Roll Command Positio_5",
        "Capt Roll Command Positio_6",
        "Capt Roll Command Positio_7",
        "Capt Roll Command Positio_8",
    ]
    wind_cols = ["Wind speed"]
    required = height_cols + pitch_cols + roll_cols + wind_cols
    if not has_columns(data, required):
        return None

    indexes = range_indexes_by_height(data, height_cols, lower, upper)
    if not indexes:
        return None

    x_values = numeric_array(data.loc[indexes], roll_cols).reshape(-1)
    y_values = -numeric_array(data.loc[indexes], pitch_cols).reshape(-1)
    wind_values = numeric_array(data.loc[indexes], wind_cols).reshape(-1)
    return control_metrics(x_values, y_values, wind_values)


def _collect_pw_phase_metrics(path, lower, upper, leap=False):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    height_cols = [
        "RADIO ALTITUDE SYS. 1_1",
        "RADIO ALTITUDE SYS. 1_2",
        "RADIO ALTITUDE SYS. 1_3",
        "RADIO ALTITUDE SYS. 1_4",
        "RADIO ALTITUDE SYS. 1_5",
    ] if leap else [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
        "RADIO HEIGHT (R/A1) SYS. 1_5",
    ]
    pitch_cols = [
        "PITCH CAPT CMD POSITION_58",
        "PITCH CAPT CMD POSITION_186",
        "PITCH CAPT CMD POSITION_314",
        "PITCH CAPT CMD POSITION_442",
        "PITCH CAPT CMD POSITION_570",
        "PITCH CAPT CMD POSITION_698",
        "PITCH CAPT CMD POSITION_826",
        "PITCH CAPT CMD POSITION_954",
    ]
    roll_cols = [
        "ROLL CAPT CMD POSITION_46",
        "ROLL CAPT CMD POSITION_174",
        "ROLL CAPT CMD POSITION_302",
        "ROLL CAPT CMD POSITION_430",
        "ROLL CAPT CMD POSITION_558",
        "ROLL CAPT CMD POSITION_686",
        "ROLL CAPT CMD POSITION_814",
        "ROLL CAPT CMD POSITION_942",
    ]
    wind_col = "WIND SPEED"
    required = height_cols + pitch_cols + roll_cols + [wind_col]
    if not has_columns(data, required):
        return None

    indexes = range_indexes_by_height(data, height_cols, lower, upper)
    if not indexes:
        return None

    x_values = numeric_array(data.loc[indexes], roll_cols).reshape(-1)
    if leap:
        x_values = -x_values
    y_values = -numeric_array(data.loc[indexes], pitch_cols).reshape(-1)
    wind_values = numeric_array(data.loc[indexes], [wind_col]).reshape(-1)
    return control_metrics(x_values, y_values, wind_values)


def collect_phase_metrics(task, lower, upper):
    collectors = {
        "CFM": lambda path: _collect_cfm_phase_metrics(path, lower, upper),
        "LEAP": lambda path: _collect_pw_phase_metrics(path, lower, upper, leap=True),
        "PW": lambda path: _collect_pw_phase_metrics(path, lower, upper, leap=False),
    }
    metric_parts = []
    for path in task.paths:
        try:
            metrics = collectors[task.aircraft](path)
        except Exception:
            continue
        if metrics:
            metric_parts.append(metrics)
    return merge_control_metrics(metric_parts, require_wind=True)


def phase_analysis_parts(task, lower, upper, phase_name):
    metrics = collect_phase_metrics(task, lower, upper)
    if metrics is None:
        return [analysis_part(f"{task.aircraft}机型{phase_name}数据无法自动计算。")]

    return [
        analysis_part(f"{task.aircraft}机型"),
        metric_analysis_part("带杆量最大", metrics["pull"], "个单位"),
        analysis_part("，"),
        metric_analysis_part("稳杆量最大", metrics["steady"], "个单位"),
        analysis_part("，"),
        metric_analysis_part("横侧量最大为", metrics["lateral"], "个单位"),
        analysis_part("，"),
        metric_analysis_part("最大侧风", metrics["wind"], "节"),
        analysis_part("。"),
    ]


def phase_analysis_text(task, lower, upper, phase_name):
    return "".join(part["text"] for part in phase_analysis_parts(task, lower, upper, phase_name))


def phase_analysis_alert(task, lower, upper):
    return control_metrics_alert(collect_phase_metrics(task, lower, upper), include_wind=True)


def _collect_cfm_throttle_values(path):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    range_height_cols = [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
        "RADIO HEIGHT (R/A1) SYS. 2_1",
        "RADIO HEIGHT (R/A1) SYS. 2_2",
        "RADIO HEIGHT (R/A1) SYS. 2_3",
        "RADIO HEIGHT (R/A1) SYS. 2_4",
    ]
    plot_height_cols = [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
    ]
    throttle_col = "Throttle lever angle Eng1_1"
    required = range_height_cols + plot_height_cols + [throttle_col]
    if not has_columns(data, required):
        return np.array([])

    indexes = range_indexes_by_height(data, range_height_cols, 0, 100)
    if not indexes:
        return np.array([])

    sorted_indexes = sorted(indexes)
    height_values = numeric_array(data.loc[sorted_indexes], plot_height_cols).reshape(-1)
    throttle_values = numeric_array(data.loc[sorted_indexes], [throttle_col]).reshape(-1)
    for index, throttle in enumerate(throttle_values):
        if throttle < 0:
            marker_pos = index * len(plot_height_cols)
            if marker_pos < len(height_values):
                return finite_values([height_values[marker_pos]])
            break
    return np.array([])


def _collect_pw_throttle_values(path, leap=False):
    data = read_qar_csv(path).drop(index=0).reset_index(drop=True)
    height_cols = [
        "RADIO ALTITUDE SYS. 1_1",
        "RADIO ALTITUDE SYS. 1_2",
        "RADIO ALTITUDE SYS. 1_3",
        "RADIO ALTITUDE SYS. 1_4",
        "RADIO ALTITUDE SYS. 1_5",
    ] if leap else [
        "RADIO HEIGHT (R/A1) SYS. 1_1",
        "RADIO HEIGHT (R/A1) SYS. 1_2",
        "RADIO HEIGHT (R/A1) SYS. 1_3",
        "RADIO HEIGHT (R/A1) SYS. 1_4",
        "RADIO HEIGHT (R/A1) SYS. 1_5",
    ]
    throttle_col = "THRUST LEVER ANGLE POS SYS. 1" if leap else "TRA (THROTTLE RESOLVER ANGLE) SYS. 1"
    ground_cols = [
        "GROUND FLIGHT BOOLEAN BSOL_26",
        "GROUND FLIGHT BOOLEAN BSOL_282",
        "GROUND FLIGHT BOOLEAN BSOL_538",
        "GROUND FLIGHT BOOLEAN BSOL_794",
    ]
    required = height_cols + [throttle_col]
    if leap:
        required += ground_cols
    if not has_columns(data, required):
        return np.array([])

    if leap:
        max_height = numeric_columns(data, height_cols).max(axis=1)
        indexes = []
        started = False
        for index in reversed(max_height.index):
            if all(data.loc[index, column] == "NOACT" for column in ground_cols):
                started = True
            if started and 0 < max_height.loc[index] < 50:
                indexes.append(index)
            elif started and max_height.loc[index] > 50:
                break
    else:
        indexes = range_indexes_by_height(data, height_cols, 0, 50)
    if not indexes:
        return np.array([])

    sorted_indexes = sorted(indexes)
    height_matrix = numeric_array(data.loc[sorted_indexes], height_cols)
    height_values = np.array([sorted(row, reverse=True) for row in height_matrix]).reshape(-1)
    throttle_values = numeric_array(data.loc[sorted_indexes], [throttle_col]).reshape(-1)
    for index, throttle in enumerate(throttle_values):
        if throttle < 0:
            marker_pos = index * (4 if leap else len(height_cols)) + (1 if leap else 0)
            if marker_pos < len(height_values):
                return finite_values([height_values[marker_pos]])
            break
    return np.array([])


def collect_throttle_values(task):
    collectors = {
        "CFM": _collect_cfm_throttle_values,
        "LEAP": lambda path: _collect_pw_throttle_values(path, leap=True),
        "PW": lambda path: _collect_pw_throttle_values(path, leap=False),
    }
    values = []
    for path in task.paths:
        try:
            values_part = collectors[task.aircraft](path)
        except Exception:
            continue
        if len(values_part):
            values.append(values_part)
    if not values:
        return np.array([])
    return np.concatenate(values)


def throttle_analysis_text(task):
    value_range = format_value_range(collect_throttle_values(task), "ft")
    if value_range is None:
        return f"{task.aircraft}机型收油门高度无法自动计算。"
    return f"{task.aircraft}机型收油门高度范围在{value_range}。"


def text_analysis_parts(text):
    return [analysis_part(text)]


def join_analysis_texts(parts):
    cleaned = [part.rstrip("。") for part in parts if part]
    if not cleaned:
        return "该项数据无法自动计算。"
    return "；".join(cleaned) + "。"


def build_analysis_entries(aircraft_tasks):
    tasks = [aircraft_tasks[aircraft] for aircraft in AIRCRAFT_ORDER if aircraft in aircraft_tasks]
    if not tasks:
        return [
            {
                "title": "分析结果：",
                "body_parts": [analysis_part("该人员无可写入的机型图片和分析数据。")],
                "alert": False,
            }
        ]

    return [
        {
            "title": "起飞50ft以下杆量：",
            "body_parts": join_analysis_parts(takeoff_analysis_parts(task) for task in tasks),
            "alert": any(takeoff_analysis_alert(task) for task in tasks),
        },
        {
            "title": "离地姿态：",
            "body_parts": join_analysis_parts(text_analysis_parts(liftoff_analysis_text(task)) for task in tasks),
            "alert": False,
        },
        {
            "title": "进近1000-100ft杆量：",
            "body_parts": join_analysis_parts(
                phase_analysis_parts(task, 100, 1000, "进近1000-100ft杆量") for task in tasks
            ),
            "alert": any(phase_analysis_alert(task, 100, 1000) for task in tasks),
        },
        {
            "title": "落地100ft以下杆量：",
            "body_parts": join_analysis_parts(
                phase_analysis_parts(task, 0, 100, "落地100ft以下杆量") for task in tasks
            ),
            "alert": any(phase_analysis_alert(task, 0, 100) for task in tasks),
        },
        {
            "title": "50ft以下曲线及收油门时机：",
            "body_parts": join_analysis_parts(text_analysis_parts(throttle_analysis_text(task)) for task in tasks),
            "alert": False,
        },
    ]


def _set_run_font(run, size=None, bold=None, color=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    latin_font = "Calibri"
    cjk_font = "Microsoft YaHei"
    run.font.name = latin_font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:cs"), latin_font)


def _set_style_font(style, size=None, bold=None, color=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    latin_font = "Calibri"
    cjk_font = "Microsoft YaHei"
    style.font.name = latin_font
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:cs"), latin_font)


def _add_paragraph(document, text="", style=None, size=11, bold=False, color=None):
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    if text:
        run = paragraph.add_run(text)
        _set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def _add_analysis_bullet(document, entry):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph(style="List Bullet")
    ppr = paragraph._p.get_or_add_pPr()
    existing_num_pr = ppr.find(qn("w:numPr"))
    if existing_num_pr is not None:
        ppr.remove(existing_num_pr)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)

    title_run = paragraph.add_run(entry["title"])
    _set_run_font(title_run, size=11, bold=True, color="2E74B5")
    body_parts = entry.get("body_parts") or [analysis_part(entry.get("body", ""), entry.get("alert", False))]
    for body_part in body_parts:
        text = body_part.get("text", "")
        if not text:
            continue
        body_run = paragraph.add_run(text)
        _set_run_font(body_run, size=11, color="C00000" if body_part.get("alert") else None)
    return paragraph


def _add_heading(document, text, level=1):
    paragraph = document.add_heading("", level=level)
    run = paragraph.add_run(text)
    _set_run_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5")
    return paragraph


def configure_report_styles(document):
    from docx.shared import Inches, Pt

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    _set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    _set_style_font(styles["Heading 1"], size=16, bold=True, color="2E74B5")
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    _set_style_font(styles["Heading 2"], size=13, bold=True, color="2E74B5")
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    if "List Bullet" in styles:
        _set_style_font(styles["List Bullet"], size=11)


def build_report_items(successful_tasks, skipped, validation_errors, failures):
    items = [
        ("title", "QAR批量分析报告"),
        (
            "paragraph",
            "本报告由批量分析程序自动生成。所有人员和不同机型写入同一个文档，按人员依次汇总 CFM、LEAP、PW 的分析图片与基础文字结论。",
        ),
    ]
    if skipped or validation_errors or failures:
        items.append(
            (
                "note",
                f"运行提示：跳过{len(skipped)}项，数据校验错误{len(validation_errors)}项，分析失败{len(failures)}项。",
            )
        )

    tasks_by_person = grouped_tasks(successful_tasks)
    for person, aircraft_tasks in tasks_by_person.items():
        items.append(("person", person))

        counts = {aircraft: len(aircraft_tasks[aircraft].paths) if aircraft in aircraft_tasks else 0 for aircraft in AIRCRAFT_ORDER}
        total = sum(counts.values())
        for aircraft in AIRCRAFT_ORDER:
            task = aircraft_tasks.get(aircraft)
            if task is None:
                continue

            items.append(("aircraft", AIRCRAFT_DISPLAY[aircraft]))
            if task.output_path.exists():
                items.append(("image", task.output_path))
            else:
                items.append(("warning", f"未找到结果图片：{task.output_path}"))

        items.append(("analysis_label", "分析："))
        items.append(
            (
                "paragraph",
                (
                    f"分析：共选取当月CFM：{counts['CFM']}段，"
                    f"LEAP：{counts['LEAP']}段，PW：{counts['PW']}段，"
                    f"共计{total}段正常航班数据。"
                ),
            )
        )
        for analysis_entry in build_analysis_entries(aircraft_tasks):
            items.append(("analysis_bullet", analysis_entry))

    return items


def generate_report(successful_tasks, skipped, validation_errors, failures):
    try:
        return generate_report_with_python_docx(successful_tasks, skipped, validation_errors, failures)
    except ModuleNotFoundError as exc:
        if exc.name != "docx":
            raise
        return generate_report_with_ooxml(successful_tasks, skipped, validation_errors, failures)


def generate_report_with_python_docx(successful_tasks, skipped, validation_errors, failures):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    report_path = unique_report_path()
    document = Document()
    configure_report_styles(document)

    for item_type, value in build_report_items(successful_tasks, skipped, validation_errors, failures):
        if item_type == "title":
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(value)
            _set_run_font(title_run, size=20, bold=True)
        elif item_type == "person":
            _add_heading(document, value, level=1)
        elif item_type == "aircraft":
            _add_heading(document, value, level=2)
        elif item_type == "image":
            document.add_picture(str(value), width=Inches(6.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif item_type == "warning":
            _add_paragraph(document, value, color="9B1C1C")
        elif item_type == "note":
            _add_paragraph(document, value, size=10, color="555555")
        elif item_type == "analysis_label":
            _add_paragraph(document, value, bold=True)
        elif item_type == "analysis_bullet":
            _add_analysis_bullet(document, value)
        else:
            _add_paragraph(document, value)

    document.save(report_path)
    return report_path


def _image_size(path):
    with open(path, "rb") as image_file:
        header = image_file.read(24)
    if header.startswith(b"\x89PNG\r\n\x1a\n") and len(header) >= 24:
        return struct.unpack(">II", header[16:24])
    return 1200, 600


def _paragraph_xml(text, style=None, bold=False, color=None, align=None):
    escaped = escape(str(text))
    ppr_parts = []
    if style:
        ppr_parts.append(f'<w:pStyle w:val="{style}"/>')
    if align:
        ppr_parts.append(f'<w:jc w:val="{align}"/>')
    ppr = f"<w:pPr>{''.join(ppr_parts)}</w:pPr>" if ppr_parts else ""
    run_props = [
        '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/>',
    ]
    if bold:
        run_props.append("<w:b/>")
    if color:
        run_props.append(f'<w:color w:val="{color}"/>')
    rpr = f"<w:rPr>{''.join(run_props)}</w:rPr>"
    return f"<w:p>{ppr}<w:r>{rpr}<w:t>{escaped}</w:t></w:r></w:p>"


def _run_xml(text, bold=False, color=None):
    escaped = escape(str(text))
    run_props = [
        '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/>',
    ]
    if bold:
        run_props.append("<w:b/>")
    if color:
        run_props.append(f'<w:color w:val="{color}"/>')
    return f"<w:r><w:rPr>{''.join(run_props)}</w:rPr><w:t>{escaped}</w:t></w:r>"


def _analysis_bullet_xml(entry):
    body_parts = entry.get("body_parts") or [analysis_part(entry.get("body", ""), entry.get("alert", False))]
    body_xml = "".join(
        _run_xml(part.get("text", ""), color="C00000" if part.get("alert") else None)
        for part in body_parts
        if part.get("text")
    )
    return (
        "<w:p>"
        "<w:pPr>"
        '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
        '<w:spacing w:after="120"/>'
        "</w:pPr>"
        f"{_run_xml(entry['title'], bold=True, color='2E74B5')}"
        f"{body_xml}"
        "</w:p>"
    )


def _image_xml(rel_id, image_name, cx, cy):
    escaped_name = escape(image_name)
    return f"""
<w:p>
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{rel_id[3:]}" name="{escaped_name}"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="{escaped_name}"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{rel_id}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{cx}" cy="{cy}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>
"""


def _styles_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults>
    <w:rPrDefault><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/><w:sz w:val="22"/></w:rPr></w:rPrDefault>
    <w:pPrDefault><w:pPr><w:spacing w:after="120"/></w:pPr></w:pPrDefault>
  </w:docDefaults>
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/>
    <w:pPr><w:spacing w:before="320" w:after="160"/><w:outlineLvl w:val="0"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/><w:b/><w:color w:val="2E74B5"/><w:sz w:val="32"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/>
    <w:pPr><w:spacing w:before="240" w:after="120"/><w:outlineLvl w:val="1"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/><w:b/><w:color w:val="2E74B5"/><w:sz w:val="26"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="ListBullet">
    <w:name w:val="List Bullet"/><w:basedOn w:val="Normal"/><w:uiPriority w:val="99"/><w:semiHidden/><w:unhideWhenUsed/>
    <w:pPr><w:numPr><w:numId w:val="1"/></w:numPr></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft YaHei" w:cs="Calibri"/><w:sz w:val="22"/></w:rPr>
  </w:style>
</w:styles>
"""


def _numbering_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="•"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>
"""


def _settings_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:w10="urn:schemas-microsoft-com:office:word"
 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
 xmlns:sl="http://schemas.openxmlformats.org/schemaLibrary/2006/main"
 mc:Ignorable="w14">
  <w:zoom w:val="bestFit"/>
  <w:defaultTabStop w:val="720"/>
  <w:characterSpacingControl w:val="doNotCompress"/>
  <w:compat>
    <w:useFELayout/>
    <w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="14"/>
  </w:compat>
  <w:themeFontLang w:val="zh-CN" w:eastAsia="zh-CN"/>
</w:settings>
"""


def _web_settings_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:allowPNG/>
</w:webSettings>
"""


def _font_table_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:font w:name="Calibri">
    <w:charset w:val="00"/>
    <w:family w:val="swiss"/>
    <w:pitch w:val="variable"/>
  </w:font>
  <w:font w:name="Microsoft YaHei">
    <w:charset w:val="86"/>
    <w:family w:val="swiss"/>
    <w:pitch w:val="variable"/>
  </w:font>
</w:fonts>
"""


def _theme_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Office Theme">
  <a:themeElements>
    <a:clrScheme name="Office">
      <a:dk1><a:sysClr val="windowText" lastClr="000000"/></a:dk1>
      <a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1>
      <a:dk2><a:srgbClr val="44546A"/></a:dk2>
      <a:lt2><a:srgbClr val="E7E6E6"/></a:lt2>
      <a:accent1><a:srgbClr val="4472C4"/></a:accent1>
      <a:accent2><a:srgbClr val="ED7D31"/></a:accent2>
      <a:accent3><a:srgbClr val="A5A5A5"/></a:accent3>
      <a:accent4><a:srgbClr val="FFC000"/></a:accent4>
      <a:accent5><a:srgbClr val="5B9BD5"/></a:accent5>
      <a:accent6><a:srgbClr val="70AD47"/></a:accent6>
      <a:hlink><a:srgbClr val="0563C1"/></a:hlink>
      <a:folHlink><a:srgbClr val="954F72"/></a:folHlink>
    </a:clrScheme>
    <a:fontScheme name="Office">
      <a:majorFont><a:latin typeface="Calibri Light"/><a:ea typeface="Microsoft YaHei"/><a:cs typeface="Calibri"/></a:majorFont>
      <a:minorFont><a:latin typeface="Calibri"/><a:ea typeface="Microsoft YaHei"/><a:cs typeface="Calibri"/></a:minorFont>
    </a:fontScheme>
    <a:fmtScheme name="Office">
      <a:fillStyleLst>
        <a:solidFill><a:schemeClr val="phClr"/></a:solidFill>
        <a:gradFill rotWithShape="1">
          <a:gsLst>
            <a:gs pos="0"><a:schemeClr val="phClr"><a:lumMod val="110000"/><a:satMod val="105000"/><a:tint val="67000"/></a:schemeClr></a:gs>
            <a:gs pos="50000"><a:schemeClr val="phClr"><a:lumMod val="105000"/><a:satMod val="103000"/><a:tint val="73000"/></a:schemeClr></a:gs>
            <a:gs pos="100000"><a:schemeClr val="phClr"><a:lumMod val="105000"/><a:satMod val="109000"/><a:tint val="81000"/></a:schemeClr></a:gs>
          </a:gsLst>
          <a:lin ang="5400000" scaled="0"/>
        </a:gradFill>
        <a:gradFill rotWithShape="1">
          <a:gsLst>
            <a:gs pos="0"><a:schemeClr val="phClr"><a:satMod val="103000"/><a:lumMod val="102000"/><a:tint val="94000"/></a:schemeClr></a:gs>
            <a:gs pos="50000"><a:schemeClr val="phClr"><a:satMod val="110000"/><a:lumMod val="100000"/><a:shade val="100000"/></a:schemeClr></a:gs>
            <a:gs pos="100000"><a:schemeClr val="phClr"><a:lumMod val="99000"/><a:satMod val="120000"/><a:shade val="78000"/></a:schemeClr></a:gs>
          </a:gsLst>
          <a:lin ang="5400000" scaled="0"/>
        </a:gradFill>
      </a:fillStyleLst>
      <a:lnStyleLst>
        <a:ln w="6350" cap="flat" cmpd="sng" algn="ctr"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:prstDash val="solid"/><a:miter lim="800000"/></a:ln>
        <a:ln w="12700" cap="flat" cmpd="sng" algn="ctr"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:prstDash val="solid"/><a:miter lim="800000"/></a:ln>
        <a:ln w="19050" cap="flat" cmpd="sng" algn="ctr"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:prstDash val="solid"/><a:miter lim="800000"/></a:ln>
      </a:lnStyleLst>
      <a:effectStyleLst>
        <a:effectStyle><a:effectLst/></a:effectStyle>
        <a:effectStyle><a:effectLst/></a:effectStyle>
        <a:effectStyle>
          <a:effectLst>
            <a:outerShdw blurRad="57150" dist="19050" dir="5400000" algn="ctr" rotWithShape="0">
              <a:srgbClr val="000000"><a:alpha val="63000"/></a:srgbClr>
            </a:outerShdw>
          </a:effectLst>
        </a:effectStyle>
      </a:effectStyleLst>
      <a:bgFillStyleLst>
        <a:solidFill><a:schemeClr val="phClr"/></a:solidFill>
        <a:gradFill rotWithShape="1">
          <a:gsLst>
            <a:gs pos="0"><a:schemeClr val="phClr"><a:tint val="40000"/><a:satMod val="350000"/></a:schemeClr></a:gs>
            <a:gs pos="40000"><a:schemeClr val="phClr"><a:tint val="45000"/><a:shade val="99000"/><a:satMod val="350000"/></a:schemeClr></a:gs>
            <a:gs pos="100000"><a:schemeClr val="phClr"><a:shade val="20000"/><a:satMod val="255000"/></a:schemeClr></a:gs>
          </a:gsLst>
          <a:path path="circle"><a:fillToRect l="50000" t="-80000" r="50000" b="180000"/></a:path>
        </a:gradFill>
        <a:gradFill rotWithShape="1">
          <a:gsLst>
            <a:gs pos="0"><a:schemeClr val="phClr"><a:tint val="80000"/><a:satMod val="300000"/></a:schemeClr></a:gs>
            <a:gs pos="100000"><a:schemeClr val="phClr"><a:shade val="30000"/><a:satMod val="200000"/></a:schemeClr></a:gs>
          </a:gsLst>
          <a:path path="circle"><a:fillToRect l="50000" t="50000" r="50000" b="50000"/></a:path>
        </a:gradFill>
      </a:bgFillStyleLst>
    </a:fmtScheme>
  </a:themeElements>
  <a:objectDefaults/>
  <a:extraClrSchemeLst/>
</a:theme>
"""


def _core_properties_xml():
    timestamp = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:dcterms="http://purl.org/dc/terms/"
 xmlns:dcmitype="http://purl.org/dc/dcmitype/"
 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>QAR批量分析报告</dc:title>
  <dc:creator>QAR Batch Analyzer</dc:creator>
  <cp:lastModifiedBy>QAR Batch Analyzer</cp:lastModifiedBy>
  <cp:revision>1</cp:revision>
  <dcterms:created xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:modified>
</cp:coreProperties>
"""


def _app_properties_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Template>Normal.dotm</Template>
  <TotalTime>0</TotalTime>
  <Pages>1</Pages>
  <Words>0</Words>
  <Characters>0</Characters>
  <Application>Microsoft Word</Application>
  <DocSecurity>0</DocSecurity>
  <Lines>0</Lines>
  <Paragraphs>0</Paragraphs>
  <ScaleCrop>false</ScaleCrop>
  <Company/>
  <LinksUpToDate>false</LinksUpToDate>
  <CharactersWithSpaces>0</CharactersWithSpaces>
  <SharedDoc>false</SharedDoc>
  <HyperlinksChanged>false</HyperlinksChanged>
  <AppVersion>16.0000</AppVersion>
</Properties>
"""


def generate_report_with_ooxml(successful_tasks, skipped, validation_errors, failures):
    report_path = unique_report_path()
    temp_dir = report_path.with_suffix(".docx_parts")
    doc_props_dir = temp_dir / "docProps"
    word_dir = temp_dir / "word"
    rels_dir = word_dir / "_rels"
    media_dir = word_dir / "media"
    theme_dir = word_dir / "theme"
    root_rels_dir = temp_dir / "_rels"
    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    doc_props_dir.mkdir(parents=True)
    media_dir.mkdir(parents=True)
    rels_dir.mkdir(parents=True)
    theme_dir.mkdir(parents=True)
    root_rels_dir.mkdir(parents=True)

    body_parts = []
    rels = []
    image_index = 1
    rel_index = 9
    image_width_emu = int(6.2 * 914400)

    for item_type, value in build_report_items(successful_tasks, skipped, validation_errors, failures):
        if item_type == "title":
            body_parts.append(_paragraph_xml(value, bold=True, align="center"))
        elif item_type == "person":
            body_parts.append(_paragraph_xml(value, style="Heading1", bold=True, color="2E74B5"))
        elif item_type == "aircraft":
            body_parts.append(_paragraph_xml(value, style="Heading2", bold=True, color="2E74B5"))
        elif item_type == "image":
            source_path = Path(value)
            image_name = f"image{image_index}{source_path.suffix.lower() or '.png'}"
            target_path = media_dir / image_name
            shutil.copyfile(source_path, target_path)
            width_px, height_px = _image_size(source_path)
            image_height_emu = int(image_width_emu * height_px / max(width_px, 1))
            rel_id = f"rId{rel_index}"
            body_parts.append(_image_xml(rel_id, image_name, image_width_emu, image_height_emu))
            rels.append(
                f'<Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{image_name}"/>'
            )
            image_index += 1
            rel_index += 1
        elif item_type == "warning":
            body_parts.append(_paragraph_xml(value, color="9B1C1C"))
        elif item_type == "note":
            body_parts.append(_paragraph_xml(value, color="555555"))
        elif item_type == "analysis_label":
            body_parts.append(_paragraph_xml(value, bold=True))
        elif item_type == "analysis_bullet":
            body_parts.append(_analysis_bullet_xml(value))
        else:
            body_parts.append(_paragraph_xml(value))

    section_props = """
<w:sectPr>
  <w:pgSz w:w="12240" w:h="15840"/>
  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/>
</w:sectPr>
"""
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
 xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
 xmlns:w10="urn:schemas-microsoft-com:office:word"
 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
 xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
 xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
 xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
 mc:Ignorable="w14">
  <w:body>
    {''.join(body_parts)}
    {section_props}
  </w:body>
</w:document>
"""

    (doc_props_dir / "core.xml").write_text(_core_properties_xml(), encoding="utf-8")
    (doc_props_dir / "app.xml").write_text(_app_properties_xml(), encoding="utf-8")
    (word_dir / "document.xml").write_text(document_xml, encoding="utf-8")
    (word_dir / "styles.xml").write_text(_styles_xml(), encoding="utf-8")
    (word_dir / "numbering.xml").write_text(_numbering_xml(), encoding="utf-8")
    (word_dir / "settings.xml").write_text(_settings_xml(), encoding="utf-8")
    (word_dir / "webSettings.xml").write_text(_web_settings_xml(), encoding="utf-8")
    (word_dir / "fontTable.xml").write_text(_font_table_xml(), encoding="utf-8")
    (theme_dir / "theme1.xml").write_text(_theme_xml(), encoding="utf-8")
    (rels_dir / "document.xml.rels").write_text(
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings" Target="webSettings.xml"/>'
        '<Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable" Target="fontTable.xml"/>'
        '<Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/>'
        '<Relationship Id="rId6" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>'
        + ''.join(rels)
        + '</Relationships>',
        encoding="utf-8",
    )
    (root_rels_dir / ".rels").write_text(
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
        '</Relationships>',
        encoding="utf-8",
    )
    (temp_dir / "[Content_Types].xml").write_text(
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Default Extension="png" ContentType="image/png"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        '<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>'
        '<Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>'
        '<Override PartName="/word/webSettings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.webSettings+xml"/>'
        '<Override PartName="/word/fontTable.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"/>'
        '<Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>'
        '</Types>',
        encoding="utf-8",
    )

    with zipfile.ZipFile(report_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for part in temp_dir.rglob("*"):
            if part.is_file():
                archive.write(part, part.relative_to(temp_dir).as_posix())
    shutil.rmtree(temp_dir)
    return report_path


class ProgressWindow:
    def __init__(self, root, total):
        self.root = root
        self.total = total
        self.done = 0

        root.title("QAR 批量分析进度")
        root.geometry("520x210")
        root.resizable(False, False)

        self.person_var = tk.StringVar(value="当前人员：等待开始")
        self.aircraft_var = tk.StringVar(value="当前机型：-")
        self.progress_var = tk.StringVar(value=f"进度：0 / {total}")
        self.status_var = tk.StringVar(value="当前状态：正在准备")

        frame = ttk.Frame(root, padding=18)
        frame.pack(fill="both", expand=True)

        ttk.Label(frame, textvariable=self.person_var).pack(anchor="w", pady=(0, 6))
        ttk.Label(frame, textvariable=self.aircraft_var).pack(anchor="w", pady=(0, 6))
        ttk.Label(frame, textvariable=self.progress_var).pack(anchor="w", pady=(0, 10))

        self.progress = ttk.Progressbar(frame, maximum=max(total, 1), mode="determinate")
        self.progress.pack(fill="x", pady=(0, 12))

        ttk.Label(frame, textvariable=self.status_var, wraplength=470).pack(anchor="w")

    def update_status(self, person=None, aircraft=None, status=None, done=None):
        if person is not None:
            self.person_var.set(f"当前人员：{person}")
        if aircraft is not None:
            self.aircraft_var.set(f"当前机型：{aircraft}")
        if status is not None:
            self.status_var.set(f"当前状态：{status}")
        if done is not None:
            self.done = done
            self.progress["value"] = done
            self.progress_var.set(f"进度：{done} / {self.total}")


def worker(tasks, modules, events, skipped, validation_errors):
    completed = 0
    failures = []
    successful_tasks = []
    total = len(tasks)

    for index, task in enumerate(tasks, start=1):
        events.put(("status", task.person, task.aircraft, "正在读取 CSV", completed))
        try:
            events.put(("status", task.person, task.aircraft, "正在生成图片", completed))
            modules[task.aircraft].more_pic(
                [str(path) for path in task.paths],
                str(task.output_path),
            )
            successful_tasks.append(task)
            completed += 1
            events.put(("status", task.person, task.aircraft, f"已生成：{task.output_path.name}", completed))
        except Exception as exc:
            failures.append((task.person, task.aircraft, str(exc), traceback.format_exc()))
            completed += 1
            events.put(("status", task.person, task.aircraft, "分析失败，已继续后续任务", completed))

    report_path = None
    report_error = None
    if successful_tasks:
        try:
            events.put(("status", "汇总报告", "DOCX", "正在生成 DOCX 分析报告", completed))
            report_path = generate_report(successful_tasks, skipped, validation_errors, failures)
            events.put(("status", "汇总报告", "DOCX", f"已生成报告：{report_path.name}", completed))
        except Exception as exc:
            report_error = (str(exc), traceback.format_exc())
            events.put(("status", "汇总报告", "DOCX", "DOCX 报告生成失败", completed))

    events.put(("done", total - len(failures), failures, report_path, report_error))


def show_initial_errors(errors):
    if not errors:
        return
    lines = [message for _, _, message in errors[:10]]
    if len(errors) > 10:
        lines.append(f"还有 {len(errors) - 10} 条错误未显示。")
    messagebox.showwarning("数据校验错误", "\n".join(lines))


def run_progress_window(tasks, skipped, validation_errors):
    root = tk.Tk()
    progress_window = ProgressWindow(root, len(tasks))
    events = queue.Queue()

    try:
        modules = load_analysis_modules()
    except Exception as exc:
        messagebox.showerror("加载失败", f"分析模块加载失败：\n{exc}")
        root.destroy()
        return

    thread = threading.Thread(target=worker, args=(tasks, modules, events, skipped, validation_errors), daemon=True)
    thread.start()

    def poll_events():
        try:
            while True:
                event = events.get_nowait()
                if event[0] == "status":
                    _, person, aircraft, status, done = event
                    progress_window.update_status(person, aircraft, status, done)
                elif event[0] == "done":
                    _, completed, failures, report_path, report_error = event
                    progress_window.update_status(status="全部任务处理完成", done=len(tasks))
                    total_errors = len(validation_errors) + len(failures) + (1 if report_error else 0)
                    summary = (
                        f"批量分析完成。\n\n"
                        f"成功：{completed}\n"
                        f"跳过：{len(skipped)}\n"
                        f"错误：{total_errors}\n\n"
                        f"图片目录：{PICS_DIR}"
                    )
                    if report_path:
                        summary += f"\n报告文件：{report_path}"
                    if failures:
                        first_failure = failures[0]
                        summary += f"\n\n首个分析失败：{first_failure[0]} {first_failure[1]}：{first_failure[2]}"
                    if report_error:
                        summary += f"\n\n报告生成失败：{report_error[0]}"
                    messagebox.showinfo("完成", summary)
                    root.destroy()
                    return
        except queue.Empty:
            pass
        root.after(150, poll_events)

    root.after(150, poll_events)
    root.mainloop()


def main():
    selector = tk.Tk()
    selector.withdraw()
    selector.update()
    selected = filedialog.askdirectory(title="请选择数据文件夹", initialdir=str(ROOT_DIR / "data"))
    selector.destroy()

    if not selected:
        return

    data_root = Path(selected)
    if not data_root.exists() or not data_root.is_dir():
        messagebox.showerror("路径错误", "选择的数据文件夹不存在。")
        return

    tasks, skipped, validation_errors = scan_data_root(data_root)
    show_initial_errors(validation_errors)

    if not tasks:
        messagebox.showinfo(
            "没有可分析数据",
            f"没有找到可分析的人员机型数据。\n\n跳过：{len(skipped)}\n错误：{len(validation_errors)}",
        )
        return

    run_progress_window(tasks, skipped, validation_errors)


if __name__ == "__main__":
    main()
